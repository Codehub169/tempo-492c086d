import os
import re
import logging
from docx import Document
from PyPDF2 import PdfReader

# Configure basic logging
logger = logging.getLogger(__name__)
if not logger.hasHandlers(): # Avoid adding multiple handlers if imported multiple times or by Flask
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Refined regex patterns
EMAIL_REGEX = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
PHONE_REGEX = r"(?:\+?\d{1,3}[-\s.]?)?(?:\(?\d{2,4}\)?[-_\s.]?){2,5}\d{2,4}" # Simplified and more robust
LINKEDIN_REGEX = r"(?:https?://)?(?:www\.)?linkedin\.com/(?:in|pub|company)/[a-zA-Z0-9_-]+/?"
GITHUB_REGEX = r"(?:https?://)?(?:www\.)?github\.com/[a-zA-Z0-9_-]+/?"

SECTION_KEYWORDS = {
    'summary': ['summary', 'objective', 'about me', 'professional profile', 'personal statement', 'overview', 'about'],
    'experience': ['experience', 'work history', 'employment history', 'professional experience', 'career summary', 'projects', 'relevant experience', 'work experience'],
    'education': ['education', 'academic background', 'qualifications', 'academic history', 'scholastic record', 'academic qualifications'],
    'skills': ['skills', 'technical skills', 'proficiencies', 'core competencies', 'technical expertise', 'technologies', 'tools', 'areas of expertise']
}

PERIOD_REGEX = re.compile(
    r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\.?(?:uary|ruary|rch|ril|y|ne|ly|ust|tember|ober|ember)?\s+\d{4}|\d{4}|\d{1,2}[/-]\d{4})" # Start Date (Month Year, Year, MM/YYYY)
    r"\s*(?:-|–|—|to|\s+until\s+)\s*" # Separator
    r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\.?(?:uary|ruary|rch|ril|y|ne|ly|ust|tember|ober|ember)?\s+\d{4}|\d{4}|\d{1,2}[/-]\d{4}|Present|Current|Ongoing)" # End Date
    r"|(\b\d{4}\b)" # Single year for education graduation
    , re.IGNORECASE
)

def _extract_text_from_pdf(file_path):
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        logger.error(f"Error reading PDF {file_path}: {e}", exc_info=True)
    return text

def _extract_text_from_docx(file_path):
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        logger.error(f"Error reading DOCX {file_path}: {e}", exc_info=True)
    return text

def _try_structure_item(text_block, item_type):
    lines = [line.strip() for line in text_block.split('\n') if line.strip()]
    if not lines or len(lines) < 1: # Need at least one line for a meaningful item
        return text_block

    item = {}
    description_lines = []
    period_line_indices = set()

    # Try to find period first
    for i, line in enumerate(lines):
        match = PERIOD_REGEX.search(line)
        if match:
            item['period'] = line.strip() # Use the whole line that contains the period
            period_line_indices.add(i)
            break # Assume one period line per item is dominant

    # Collect non-period lines for title, org, description
    content_lines = [lines[i] for i in range(len(lines)) if i not in period_line_indices]

    if not content_lines: # If all lines were period lines or no content lines left
        if item.get('period'): # If we only have a period, it's not enough structured info
            return text_block
        return text_block

    # Heuristic: first content line is title/degree, second is company/institution
    if item_type == 'experience':
        item['title'] = content_lines[0]
        if len(content_lines) > 1:
            item['company'] = content_lines[1]
            description_lines = content_lines[2:]
        else:
            # If only one content line, it might be title, but company is missing. Fallback.
            return text_block 
    elif item_type == 'education':
        item['degree'] = content_lines[0]
        if len(content_lines) > 1:
            item['institution'] = content_lines[1]
            description_lines = content_lines[2:]
        else:
            # If only one content line, it might be degree, but institution is missing. Fallback.
            return text_block

    item['description'] = '\n'.join(description_lines).strip()

    # Validate essential fields for structured item
    if item_type == 'experience' and not (item.get('title') and item.get('company')):
        return text_block
    if item_type == 'education' and not (item.get('degree') and item.get('institution')):
        return text_block
    
    # Remove empty description if it's just an empty string
    if not item['description']:
        del item['description']

    return item

def _process_section_content(section_name, content_lines, target_dict):
    if not section_name or not content_lines or section_name not in target_dict:
        return

    full_content_block = '\n'.join(content_lines).strip()
    if not full_content_block:
        return

    if section_name in ['experience', 'education']:
        # Append the block; structuring will happen later
        target_dict[section_name].append(full_content_block)
    elif section_name == 'skills':
        # Split skills by common delimiters (comma, newline, semicolon, bullet points)
        raw_skills = re.split(r'[\n,;•*-]', full_content_block)
        for skill in raw_skills:
            s = skill.strip()
            if s and len(s) > 1: # Basic validation for skill length
                target_dict[section_name].append(s)
    else: # For 'summary' or other single-block text sections
        if target_dict.get(section_name):
            target_dict[section_name] += "\n" + full_content_block
        else:
            target_dict[section_name] = full_content_block

def _parse_sections(text_content):
    parsed_data = {
        'name': '', 'title': '',
        'email': '', 'phone': '', 'linkedin': '', 'github': '',
        'summary': '',
        'experience': [], 'education': [], 'skills': []
    }

    # Extract contact info
    emails = sorted(list(set(re.findall(EMAIL_REGEX, text_content, re.IGNORECASE))))
    if emails: parsed_data['email'] = emails[0] # Take the first one for simplicity
    
    phones = sorted(list(set(re.findall(PHONE_REGEX, text_content))))
    if phones: parsed_data['phone'] = phones[0]

    linkedin_urls = sorted(list(set(re.findall(LINKEDIN_REGEX, text_content, re.IGNORECASE))))
    if linkedin_urls: parsed_data['linkedin'] = linkedin_urls[0].replace("https://", "").replace("http://", "")

    github_urls = sorted(list(set(re.findall(GITHUB_REGEX, text_content, re.IGNORECASE))))
    if github_urls: parsed_data['github'] = github_urls[0].replace("https://", "").replace("http://", "")

    lines = [line.strip() for line in text_content.split('\n') if line.strip()]
    current_section = None
    temp_content = []
    
    # Attempt to get name and title (heuristic)
    # Name is often the first prominent line, title might be second or near contact details.
    if lines:
        # First line as potential name, if it's not an email or phone or too long
        first_line = lines[0]
        if not re.search(EMAIL_REGEX + "|" + PHONE_REGEX, first_line, re.IGNORECASE) and len(first_line.split()) < 6 and len(first_line) < 50:
            parsed_data['name'] = first_line
            # Try second line as title
            if len(lines) > 1:
                second_line = lines[1]
                if not re.search(EMAIL_REGEX + "|" + PHONE_REGEX, second_line, re.IGNORECASE) and len(second_line.split()) < 10 and len(second_line) < 70:
                    # Check if it looks like a section header
                    is_section_header = any(keyword in second_line.lower() for section_keywords in SECTION_KEYWORDS.values() for keyword in section_keywords)
                    if not is_section_header:
                        parsed_data['title'] = second_line
    
    # Default if not found
    if not parsed_data['name']: parsed_data['name'] = 'Your Name'
    if not parsed_data['title']: parsed_data['title'] = 'Professional Title'

    section_line_indices = set()

    for i, line in enumerate(lines):
        line_lower = line.lower()
        identified_new_section = False

        # Check if line is a section header
        for section_key_candidate, keywords in SECTION_KEYWORDS.items():
            # Section headers are usually short and contain keywords
            if any(keyword in line_lower for keyword in keywords) and len(line.split()) < 6:
                _process_section_content(current_section, temp_content, parsed_data)
                current_section = section_key_candidate
                temp_content = []
                identified_new_section = True
                section_line_indices.add(i)
                break
            if identified_new_section:
                break
        
        if not identified_new_section and i not in section_line_indices:
            # If line is not part of name/title extraction and not a section header
            if (parsed_data['name'] and line == parsed_data['name']) or \
               (parsed_data['title'] and line == parsed_data['title']): # Avoid re-adding name/title
                continue

            if current_section:
                temp_content.append(line)
            elif not parsed_data['summary'] and len(line.split()) > 3: # Content before any explicit section could be summary
                 # Avoid contact info being part of summary
                if not re.search(EMAIL_REGEX + "|" + PHONE_REGEX + "|" + LINKEDIN_REGEX + "|" + GITHUB_REGEX, line, re.IGNORECASE):
                    parsed_data['summary'] += line + "\n"

    _process_section_content(current_section, temp_content, parsed_data)

    # Post-process experience and education for structure
    parsed_data['experience'] = [(_try_structure_item(block, 'experience')) for block in parsed_data.get('experience', [])]
    parsed_data['education'] = [(_try_structure_item(block, 'education')) for block in parsed_data.get('education', [])]

    # Clean up skills: unique, sensible length
    if parsed_data['skills']:
        unique_skills = list(set(s.strip().capitalize() for s in parsed_data['skills'] if s.strip() and 1 < len(s.strip()) < 50))
        parsed_data['skills'] = sorted(unique_skills)
    else:
        parsed_data['skills'] = []

    if parsed_data['summary']:
        parsed_data['summary'] = parsed_data['summary'].strip()

    return parsed_data

def parse_resume(file_path):
    _, file_extension = os.path.splitext(file_path)
    text_content = ""
    ext_lower = file_extension.lower()

    if ext_lower == '.pdf':
        text_content = _extract_text_from_pdf(file_path)
    elif ext_lower in ['.docx', '.doc']:
        text_content = _extract_text_from_docx(file_path)
    else:
        allowed_types = "PDF, DOCX, DOC"
        logger.warning(f"Unsupported file type: {ext_lower}. File: {file_path}")
        raise ValueError(f"Unsupported file type: {ext_lower}. Only {allowed_types} are supported.")

    if not text_content.strip():
        logger.warning(f"No text extracted from {file_path}. Document might be image-based or empty.")
        return {
            'name': 'Error: Could Not Parse Name',
            'title': 'Error: Could Not Parse Title',
            'email': '', 'phone': '', 'linkedin': '', 'github': '',
            'summary': 'Could not extract text from the resume. The document might be image-based, corrupted, or empty.',
            'experience': [], 'education': [], 'skills': []
        }

    parsed_data = _parse_sections(text_content)
    logger.info(f"Successfully parsed resume: {file_path}")
    return parsed_data

if __name__ == '__main__':
    # Example Usage (for testing)
    # Create a dummy docx for testing
    dummy_docx_path = 'dummy_resume_parser_test.docx'
    try:
        doc = Document()
        doc.add_heading('Dr. Jane Doe', 0)
        doc.add_paragraph('Senior Quantum Physicist & Bagel Enthusiast')
        doc.add_paragraph('jane.doe@example.com | (555) 123-4567 | linkedin.com/in/janedoe | github.com/janedoe')
        doc.add_heading('Overview', level=1)
        doc.add_paragraph('A highly motivated physicist with a passion for bagels.')
        doc.add_heading('Professional Experience', level=1)
        doc.add_paragraph('Lead Physicist\nQuantum Bagel Labs\nJan 2020 - Present\n- Spearheaded research on bagel quantum properties.')
        doc.add_paragraph('Research Fellow\nInstitute of Theoretical Breakfast\n2018 - 2019\n- Studied Schr\u00f6dinger\'s cat breakfast preferences.')
        doc.add_heading('Education', level=1)
        doc.add_paragraph('Ph.D. in Quantum Bagel Dynamics\nUniversity of Advanced Studies\nGraduated: May 2015')
        doc.add_paragraph('B.S. in Physics\nState College\n2011')        
        doc.add_heading('Technical Skills', level=1)
        doc.add_paragraph('Quantum Entanglement, Python, LaTeX, Bagel Slicing; Cream Cheese Spreading')
        doc.save(dummy_docx_path)
        logger.info(f"Dummy DOCX created at {dummy_docx_path}")
        
        parsed_docx_data = parse_resume(dummy_docx_path)
        logger.info("\n--- Parsed DOCX Data ---")
        import json
        logger.info(json.dumps(parsed_docx_data, indent=2))

    except Exception as e:
        logger.error(f"Error in DOCX example processing: {e}", exc_info=True)
    finally:
        if os.path.exists(dummy_docx_path):
            os.remove(dummy_docx_path)
            logger.info(f"Removed dummy DOCX: {dummy_docx_path}")
